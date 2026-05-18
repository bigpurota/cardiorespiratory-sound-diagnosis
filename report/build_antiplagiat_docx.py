#!/usr/bin/env python
"""Build report/Report_Tsember_antiplagiat.docx from the Typst sources.

A faithful PLAIN-TEXT rendering of the report for the Antiplagiat self-check:
- strips Typst markup, renders inline math as readable text (MAcc = 0.903, ±, ≈ …),
- keeps the heading hierarchy (Word Heading 1/2/3) so the structure panel detects
  Annotation / Introduction / Bibliography / Annexes,
- resolves [n] citations in order of first appearance and prints the bibliography,
- Annex-7 page setup (TNR 12, 1.5 spacing, justified, margins, page numbers).

Run:  uv run --with python-docx python report/build_antiplagiat_docx.py
"""
import re, pathlib

SECT = pathlib.Path("report/sections")
ORDER = ["00-abstract","01-introduction","02-ch1-litreview","03-ch2-methods",
         "04-ch3-results","05-ch4-novelty","06-conclusion","08-annexes"]

# ---- bibliography keys + citation order ------------------------------------
bib_text = pathlib.Path("report/refs.bib").read_text(encoding="utf-8")
bib_keys = re.findall(r"@\w+\{([^,]+),", bib_text)
bib_keyset = set(bib_keys)
cite_order = []   # bib keys in order of first appearance

def render_bib_entry(key):
    """Crude IEEE-ish one-liner from a .bib entry."""
    m = re.search(r"@\w+\{" + re.escape(key) + r",(.*?)\n\}", bib_text, re.S)
    body = m.group(1) if m else ""
    def f(field):
        mm = re.search(field + r"\s*=\s*[{\"](.+?)[}\"]\s*,?", body, re.S)
        return re.sub(r"\s+", " ", mm.group(1)).strip() if mm else ""
    auth, title, jour = f("author"), f("title"), f("journal") or f("booktitle") or f("publisher")
    year, vol, pages = f("year"), f("volume"), f("pages")
    doi, url = f("doi"), f("url")
    parts = [p for p in [auth, f'"{title},"' if title else "", jour,
                         f"vol. {vol}" if vol else "", f"pp. {pages}" if pages else "",
                         year, f"doi: {doi}" if doi else (url if url else "")] if p]
    return ", ".join(parts) + "."

# ---- inline Typst -> readable text -----------------------------------------
def conv_math(s):
    # s is the content between $...$
    s = s.replace("plus.minus", "±").replace("gt.eq", "≥").replace("lt.eq", "≤")
    s = s.replace("approx", "≈").replace("times", "×").replace("arrow.r", "→")
    s = s.replace("rho", "ρ").replace("\\/", "/").replace("thin", " ")
    s = re.sub(r'"([^"]*)"', r"\1", s)          # "MAcc" -> MAcc
    s = re.sub(r'_"?([A-Za-z0-9]+)"?', r"_\1", s)  # subscript Se_crk
    s = re.sub(r"\bbold\(([^)]*)\)", r"\1", s)
    s = s.replace("$", "")
    return re.sub(r"\s+", " ", s).strip()

def inline(t):
    t = re.sub(r"\$([^$]*)\$", lambda m: conv_math(m.group(1)), t)   # math
    # citations / cross-refs
    def cite(m):
        k = m.group(1)
        if k.startswith(("fig-","tab-","eq-")):
            return {"f":"the figure","t":"the table","e":"the equation"}[k[0]]
        if k in bib_keyset:
            if k not in cite_order: cite_order.append(k)
            return f"[{cite_order.index(k)+1}]"
        return ""
    t = re.sub(r"@([A-Za-z0-9_\-]+)", cite, t)
    t = re.sub(r"`([^`]*)`", r"\1", t)                 # raw code
    t = re.sub(r"\*([^*]+)\*", r"\1", t)               # bold
    t = re.sub(r"_([^_]+)_", r"\1", t)                 # italic
    t = re.sub(r"#h\([^)]*\)", " ", t)
    t = t.replace("\\#", "#").replace("\\@", "@")
    return re.sub(r"[ \t]+", " ", t).strip()

# ---- caption extraction (bracket-balanced) ---------------------------------
def extract_caption(block):
    i = block.find("caption:")
    if i < 0: return None
    j = block.find("[", i)
    if j < 0: return None
    depth, k = 0, j
    while k < len(block):
        if block[k] == "[": depth += 1
        elif block[k] == "]":
            depth -= 1
            if depth == 0: break
        k += 1
    return inline(block[j+1:k])

# ---- parse one section file into (kind, level, text) blocks ----------------
def parse_section(name, annex_letter=[0]):
    text = pathlib.Path(SECT / f"{name}.typ").read_text(encoding="utf-8")
    out = []
    lines = text.splitlines()
    i = 0
    para = []
    li = []                                            # current list item (may wrap)
    def flush():
        if para:
            s = inline(" ".join(para)).strip()
            if s: out.append(("p", 0, s))
            para.clear()
    def flush_li():
        if li:
            s = inline(" ".join(li)).strip()
            if s: out.append(("li", 0, s))
            li.clear()
    while i < len(lines):
        ln = lines[i]
        st = ln.strip()
        if st.startswith("//") or st == "":
            flush(); flush_li();
            i += 1; continue
        if st.startswith("```"):                       # code fence
            flush(); flush_li(); i += 1; code = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code.append(lines[i]); i += 1
            i += 1
            out.append(("code", 0, "\n".join(code)))
            continue
        if st.startswith(("#import","#set","#show","#counter","#state","#pagebreak","#v(")):
            flush(); flush_li(); i += 1; continue
        if st.startswith("#figure("):                  # figure/table block
            flush(); flush_li(); depth = 0; buf = []
            while i < len(lines):
                buf.append(lines[i]); depth += lines[i].count("(") - lines[i].count(")")
                i += 1
                if depth <= 0: break
            cap = extract_caption("\n".join(buf))
            if cap: out.append(("cap", 0, cap))
            continue
        m = re.match(r"^(=+)\s+(.*)$", st)              # = / == / === heading
        if m:
            flush(); flush_li(); lvl = len(m.group(1)); title = inline(m.group(2))
            if name == "08-annexes" and lvl == 1:
                annex_letter[0] += 1
                title = f"Annex {chr(64+annex_letter[0])}: {title}"
            out.append(("h", lvl, title)); i += 1; continue
        m = re.match(r"^#heading\((.*?)\)\[(.*)\]\s*$", st)   # #heading(...)[Title]
        if m:
            flush(); flush_li(); args, title = m.group(1), inline(m.group(2))
            lvl = 2 if "level: 2" in args or "level: 3" in args else 1
            out.append(("h", lvl, title)); i += 1; continue
        mli = re.match(r"^[+\-]\s+(.*)$", st)           # "+ item" / "- item" list marker
        if mli:
            flush(); flush_li()                         # close any open para/item
            li.append(mli.group(1)); i += 1; continue
        if li:                                          # wrapped continuation of a list item
            li.append(st); i += 1; continue
        para.append(st); i += 1
    flush(); flush_li()
    return out

blocks = []
for n in ORDER:
    blocks += parse_section(n)
    if n == "06-conclusion":
        blocks.append(("bib_marker", 0, ""))   # bibliography goes after conclusion

# ---- build the .docx --------------------------------------------------------
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
st = doc.styles["Normal"]; st.font.name = "Times New Roman"; st.font.size = Pt(12)
st.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
pf = st.paragraph_format; pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; pf.first_line_indent = Cm(1.25)
for h, sz in [("Heading 1",14),("Heading 2",13),("Heading 3",12)]:
    s = doc.styles[h]; s.font.name = "Times New Roman"; s.font.size = Pt(sz)
    s.font.bold = True; s.font.color.rgb = None
sec = doc.sections[0]
sec.left_margin, sec.right_margin = Cm(2.5), Cm(1.0)
sec.top_margin, sec.bottom_margin = Cm(2.0), Cm(2.0)
# page number footer (centre)
fp = sec.footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = fp.add_run()
for el, attr in [("begin",None),("instrText","PAGE"),("end",None)]:
    e = OxmlElement(f"w:fld" if False else "w:r");
fld1 = OxmlElement("w:fldChar"); fld1.set(qn("w:fldCharType"),"begin")
instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"),"preserve"); instr.text="PAGE"
fld2 = OxmlElement("w:fldChar"); fld2.set(qn("w:fldCharType"),"end")
run._r.append(fld1); run._r.append(instr); run._r.append(fld2)

# title page
def center(txt, bold=False, size=None):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run(txt); r.bold = bold
    if size: r.font.size = Pt(size)
    return p
center("NATIONAL RESEARCH UNIVERSITY HIGHER SCHOOL OF ECONOMICS", size=14)
center("Faculty of Computer Science")
center('Bachelor’s Programme "Data Science and Business Analytics"')
doc.add_paragraph("UDC 004.85:616.12").paragraph_format.first_line_indent = Cm(0)
center("Research Project Report on the Topic:")
center("Diagnosis of Pathologies of the Cardiorespiratory System and Main Arteries "
       "Based on the Analysis of Sound Data", bold=True, size=16)
doc.add_paragraph("Submitted by the Student: Group БПАД244, 2nd year of study — "
                  "Tsember Andrei Alekseevich").paragraph_format.first_line_indent = Cm(0)
doc.add_paragraph("Approved by the Project Supervisor: Tomashchuk Kornei Kirillovich, "
                  "Lecturer, Faculty of Computer Science, HSE University").paragraph_format.first_line_indent = Cm(0)
center("Moscow 2026")
doc.add_page_break()

def add_para(text, style=None, mono=False):
    p = doc.add_paragraph(style=style)
    if style and style.startswith("Heading"):
        p.paragraph_format.first_line_indent = Cm(0)
    if mono:
        p.paragraph_format.first_line_indent = Cm(0)
        r = p.add_run(text); r.font.name = "Courier New"; r.font.size = Pt(10)
    else:
        p.add_run(text)
    return p

def emit_bibliography():
    add_para("Bibliography", style="Heading 1")
    for i, k in enumerate(cite_order, 1):
        p = doc.add_paragraph(); p.paragraph_format.first_line_indent = Cm(0)
        p.add_run(f"[{i}] {render_bib_entry(k)}")

for kind, lvl, txt in blocks:
    if kind == "h": add_para(txt, style=f"Heading {min(lvl,3)}")
    elif kind == "p": add_para(txt)
    elif kind == "li":                           # bulleted list item
        p = doc.add_paragraph(f"• {txt}")
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(0.75)
    elif kind == "cap": add_para(txt)            # figure/table caption as prose
    elif kind == "code": add_para(txt, mono=True)
    elif kind == "bib_marker": emit_bibliography()

out = "report/Report_Tsember_antiplagiat.docx"
doc.save(out)
words = sum(len(t.split()) for k,l,t in blocks if k in ("p","cap","li"))
print(f"WROTE {out}")
print(f"headings: {sum(1 for k,l,t in blocks if k=='h')}  paragraphs: {sum(1 for k,l,t in blocks if k=='p')}  "
      f"list-items: {sum(1 for k,l,t in blocks if k=='li')}  "
      f"captions: {sum(1 for k,l,t in blocks if k=='cap')}  citations: {len(cite_order)}  body words≈{words}")
